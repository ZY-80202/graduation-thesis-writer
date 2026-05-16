from pathlib import Path

from docx import Document

from thesis_skill.parsers.docx_parser import read_docx_paragraphs


def test_read_docx_paragraphs(tmp_path: Path) -> None:
    path = tmp_path / "sample.docx"
    doc = Document()
    doc.add_heading("第1章 绪论", level=1)
    doc.add_paragraph("正文内容")
    doc.save(path)

    paragraphs = read_docx_paragraphs(path)

    assert paragraphs[0]["text"] == "第1章 绪论"
    assert paragraphs[1]["text"] == "正文内容"
