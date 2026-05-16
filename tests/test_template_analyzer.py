from pathlib import Path

from docx import Document

from thesis_skill.analyzer.template_analyzer import analyze_template


def test_analyze_template_extracts_regions(tmp_path: Path) -> None:
    template = tmp_path / "template.docx"
    out = tmp_path / "profiles"
    doc = Document()
    doc.add_paragraph("毕业设计说明书")
    doc.add_paragraph("摘要")
    doc.add_paragraph("目录")
    doc.add_heading("第1章 绪论", level=1)
    doc.add_paragraph("参考文献")
    doc.save(template)

    profile = analyze_template(template, out)

    assert (out / "template_profile.json").exists()
    assert "摘要" in profile.detected_keywords
    assert any(region.name == "toc" for region in profile.regions)
