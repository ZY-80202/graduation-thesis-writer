from __future__ import annotations

import re
from pathlib import Path
from typing import Any, Dict, Iterable, List, Set

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from thesis_skill.docx_writer.cover_cloner import build_cover_replacements, clone_template_cover
from thesis_skill.docx_writer.figure_manager import insert_figure_with_caption
from thesis_skill.docx_writer.numbering_manager import BACK_MATTER_TITLES, FRONT_MATTER_TITLES, heading_text
from thesis_skill.docx_writer.section_manager import configure_cover_section, create_body_section, create_front_matter_section
from thesis_skill.docx_writer.style_manager import apply_document_styles, body_style, caption_style, heading_style
from thesis_skill.docx_writer.table_manager import insert_table_with_caption
from thesis_skill.docx_writer.template_clone import prepare_template_docx
from thesis_skill.docx_writer.toc_manager import build_static_toc_with_dot_leaders, insert_toc_field
from thesis_skill.generator.caption_generator import CaptionGenerator
from thesis_skill.generator.table_generator import (
    build_database_field_tables,
    build_database_overview_table,
    build_endpoint_table,
    build_environment_table,
    build_frontend_page_table,
    build_test_table,
)
from thesis_skill.models import DiagramArtifact, ProjectProfile, ScreenshotAsset, SectionDraft, ThesisDocument
from thesis_skill.utils.file_utils import ensure_dir


def build_docx(
    template_path: str | Path,
    thesis: ThesisDocument,
    project: ProjectProfile,
    diagrams: List[DiagramArtifact] | None,
    out_path: str | Path = "outputs/final_thesis.docx",
    config: Dict[str, Any] | None = None,
    *,
    strict_template: bool = False,
    cover_pages: int = 2,
    toc_mode: str = "field",
    body_start_title: str = "概述",
    max_pages: int = 35,
) -> Path:
    """Build the Word thesis.

    In strict-template mode the output starts as a clone of the template DOCX.
    The cloned front pages keep their original OOXML layout and only field text
    is replaced. New thesis content is appended after a fresh front-matter/body
    section split.
    """

    template = Path(template_path)
    if not template.exists():
        raise FileNotFoundError(f"学校模板不存在: {template}")

    output = Path(out_path)
    ensure_dir(output.parent)
    prepared_template = prepare_template_docx(template)

    if strict_template:
        base_path = clone_template_cover(
            prepared_template,
            output,
            build_cover_replacements(config or {}),
            cover_pages=cover_pages,
        )
        document = Document(str(base_path))
        configure_cover_section(document)
    else:
        document = Document(str(prepared_template))
        _clear_body(document)
        _set_default_section(document, config)
        _add_simple_cover(document, thesis)

    apply_document_styles(document, config, template_path=prepared_template)

    create_front_matter_section(document)
    _add_abstract(document, thesis)
    _add_toc(document, thesis.sections, toc_mode=toc_mode)

    short_title = _short_title(thesis.title or body_start_title)
    create_body_section(document, short_title=short_title)

    diagram_map = {diagram.key: diagram for diagram in diagrams or []}
    inserted_diagrams: Set[str] = set()
    inserted_screenshots: Set[str] = set()
    captions = CaptionGenerator()

    for section in thesis.sections:
        _add_section(document, section, project, diagram_map, inserted_diagrams, inserted_screenshots, captions)

    _add_references(document, thesis)
    _add_acknowledgements(document, thesis)
    _add_appendices(document, thesis)
    _add_generation_notes(document, max_pages)

    document.save(str(output))
    return output


def _clear_body(document) -> None:
    body = document._element.body
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def _set_default_section(document, config: Dict[str, Any] | None) -> None:
    fmt = (config or {}).get("format", {}) or (config or {}).get("school_format", {})
    for section in document.sections:
        section.top_margin = Cm(float(fmt.get("top_margin_cm", 2.54)))
        section.bottom_margin = Cm(float(fmt.get("bottom_margin_cm", 2.54)))
        section.left_margin = Cm(float(fmt.get("left_margin_cm", 3.0)))
        section.right_margin = Cm(float(fmt.get("right_margin_cm", 2.6)))


def _add_simple_cover(document, thesis: ThesisDocument) -> None:
    for _ in range(4):
        document.add_paragraph("")
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(thesis.title or "毕业设计（论文）")
    run.bold = True
    run.font.size = Pt(22)
    for _ in range(4):
        document.add_paragraph("")
    rows = [
        ("学生姓名", thesis.author or "【请填写作者姓名】"),
        ("学号", thesis.student_id or "【请填写学号】"),
        ("学院", thesis.college or "【请填写学院】"),
        ("专业", thesis.major or "【请填写专业】"),
        ("指导教师", thesis.supervisor or "【请填写指导教师】"),
    ]
    for label, value in rows:
        paragraph = document.add_paragraph(style=body_style(document))
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run(f"{label}: {value}")
    document.add_page_break()


def _add_abstract(document, thesis: ThesisDocument) -> None:
    _front_heading(document, "摘    要", include_in_toc=True)
    _body_paragraph(document, thesis.abstract_cn or "【请补充中文摘要】")
    _body_paragraph(document, "关键词: " + "；".join(thesis.keywords or ["【请补充关键词】"]))
    _front_heading(document, "Abstract", include_in_toc=False)
    _body_paragraph(document, thesis.abstract_en or "【Please complete the English abstract.】")
    _body_paragraph(document, "Key words: " + "; ".join(thesis.keywords or ["keyword"]))
    document.add_page_break()


def _add_toc(document, sections: Iterable[SectionDraft], toc_mode: str = "field") -> None:
    _front_heading(document, "目    录", include_in_toc=False)
    if toc_mode == "static":
        build_static_toc_with_dot_leaders(document, sections, include_abstract=True)
    else:
        paragraph = document.add_paragraph()
        insert_toc_field(paragraph)
        hint = document.add_paragraph(style=body_style(document))
        hint.text = "提示: 首次打开文档后，在 Word 中右键目录并选择“更新域”即可刷新页码。"
    document.add_page_break()


def _add_section(
    document,
    section: SectionDraft,
    project: ProjectProfile,
    diagram_map: Dict[str, DiagramArtifact],
    inserted_diagrams: Set[str],
    inserted_screenshots: Set[str],
    captions: CaptionGenerator,
) -> None:
    title = heading_text(section.number, section.title)
    if section.level == 1 and section.title in BACK_MATTER_TITLES:
        _front_heading(document, section.title, include_in_toc=True)
    else:
        _heading(document, title, min(max(section.level, 1), 3))

    for paragraph in section.paragraphs:
        _body_paragraph(document, paragraph)

    _insert_section_assets(document, section, project, diagram_map, inserted_diagrams, inserted_screenshots, captions)

    for child in section.children:
        _add_section(document, child, project, diagram_map, inserted_diagrams, inserted_screenshots, captions)


def _insert_section_assets(
    document,
    section: SectionDraft,
    project: ProjectProfile,
    diagram_map: Dict[str, DiagramArtifact],
    inserted_diagrams: Set[str],
    inserted_screenshots: Set[str],
    captions: CaptionGenerator,
) -> None:
    title = section.title
    chapter = _chapter_number(section)
    asset_key = _diagram_key_for_section(title)
    if asset_key and asset_key in diagram_map and asset_key not in inserted_diagrams:
        diagram = diagram_map[asset_key]
        caption = captions.figure(chapter, diagram.title)
        insert_figure_with_caption(document, diagram.png_path, caption)
        _reference_line(document, caption)
        inserted_diagrams.add(asset_key)

    _insert_screenshots_for_section(document, section, project, inserted_screenshots, captions)

    if _has_any(title, ["开发环境", "运行环境"]):
        insert_table_with_caption(document, captions.table(chapter, "系统开发环境表"), build_environment_table(project))
    elif _has_any(title, ["功能结构", "功能模块", "页面"]):
        insert_table_with_caption(document, captions.table(chapter, "系统功能页面说明"), build_frontend_page_table(project))
    elif _has_any(title, ["数据库设计", "数据表"]):
        _add_database_design_intro(document, project)
        insert_table_with_caption(document, captions.table(chapter, "数据库表汇总"), build_database_overview_table(project))
        for caption_text, rows in build_database_field_tables(project):
            insert_table_with_caption(document, captions.table(chapter, caption_text), rows)
    elif _has_any(title, ["接口设计", "接口"]):
        insert_table_with_caption(document, captions.table(chapter, "核心接口设计"), build_endpoint_table(project))
    elif _has_any(title, ["功能测试", "测试用例"]):
        insert_table_with_caption(document, captions.table(chapter, "功能测试用例表"), build_test_table(project))


def _add_database_design_intro(document, project: ProjectProfile) -> None:
    if not project.database_tables:
        _body_paragraph(document, "项目资料中未识别到完整 database.sql、init.sql 或 schema.sql，本节保留数据库表结构待补充项。")
        return
    table_names = "、".join(table.name for table in project.database_tables[:8])
    _body_paragraph(
        document,
        f"数据库设计根据项目 SQL 脚本提取，主要包含{table_names}等数据表。表结构围绕用户身份、业务数据、操作记录和系统配置等信息展开，字段类型、主键和可空约束以 SQL 定义为准。",
    )


def _diagram_key_for_section(title: str) -> str | None:
    if _has_any(title, ["业务流程"]):
        return "business_flow"
    if _has_any(title, ["架构", "总体架构", "系统架构"]):
        return "architecture"
    if _has_any(title, ["功能模块", "功能结构"]):
        return "function_structure"
    if _has_any(title, ["数据库"]):
        return "er"
    if _has_any(title, ["接口"]):
        return "frontend_backend"
    if _has_any(title, ["测试方法", "测试流程"]):
        return "test_flow"
    if _has_any(title, ["后台", "管理"]):
        return "admin_flow"
    if _has_any(title, ["用户", "登录", "注册"]):
        return "user_flow"
    return None


def _insert_screenshots_for_section(
    document,
    section: SectionDraft,
    project: ProjectProfile,
    inserted: Set[str],
    captions: CaptionGenerator,
) -> None:
    if not project.screenshot_assets:
        return
    chapter = _chapter_number(section)
    section_text = f"{section.number} {section.title}"
    for asset in project.screenshot_assets:
        if asset.path in inserted or not _screenshot_matches_section(asset, section_text):
            continue
        image_path = Path(project.project_path) / asset.path
        caption = captions.figure(chapter, asset.caption or "系统页面运行效果")
        insert_figure_with_caption(document, image_path, caption)
        _reference_line(document, caption)
        inserted.add(asset.path)


def _screenshot_matches_section(asset: ScreenshotAsset, section_text: str) -> bool:
    target = f"{asset.inferred_section} {asset.matched_module}".strip()
    if not target:
        return False
    if asset.inferred_section and asset.inferred_section in section_text:
        return True
    module_key = (asset.matched_module or "").replace("模块", "")
    if module_key and module_key in section_text:
        return True
    if "测试" in asset.inferred_section and "测试" in section_text:
        return True
    if "数据库" in asset.inferred_section and "数据库" in section_text:
        return True
    return False


def _chapter_number(section: SectionDraft) -> int:
    match = re.match(r"(\d+)", section.number or "")
    if match:
        return int(match.group(1))
    for parent in [section.title, section.number]:
        match = re.search(r"第\s*(\d+)\s*章", parent or "")
        if match:
            return int(match.group(1))
    return 1


def _add_references(document, thesis: ThesisDocument) -> None:
    _front_heading(document, "参考文献", include_in_toc=True)
    if not thesis.references:
        _body_paragraph(document, "【请补充参考文献】")
        return
    for index, reference in enumerate(thesis.references, start=1):
        _body_paragraph(document, f"[{index}] {reference}")


def _add_acknowledgements(document, thesis: ThesisDocument) -> None:
    _front_heading(document, "致    谢", include_in_toc=True)
    _body_paragraph(document, thesis.acknowledgements or "【请补充致谢内容】")


def _add_appendices(document, thesis: ThesisDocument) -> None:
    if not thesis.appendices:
        return
    _front_heading(document, "附    录", include_in_toc=True)
    for appendix in thesis.appendices:
        _body_paragraph(document, appendix)


def _add_generation_notes(document, max_pages: int) -> None:
    paragraph = document.add_paragraph(style=caption_style(document))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.text = f"生成说明: 文档按 {max_pages} 页以内目标控制篇幅，最终页数请以 Word 更新域后的结果为准。"


def _heading(document, text: str, level: int) -> None:
    paragraph = document.add_paragraph(text, style=heading_style(document, level))
    if level == 1:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _front_heading(document, text: str, include_in_toc: bool) -> None:
    paragraph = document.add_paragraph(style=body_style(document))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    if include_in_toc:
        _set_outline_level(paragraph, 0)


def _body_paragraph(document, text: str) -> None:
    paragraph = document.add_paragraph(str(text), style=body_style(document))
    paragraph.paragraph_format.first_line_indent = Cm(0.74)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def _reference_line(document, caption: str) -> None:
    match = re.search(r"(图\s*\d+\.\d+)", caption)
    if match:
        _body_paragraph(document, f"相关页面或结构如{match.group(1)} 所示。")


def _set_outline_level(paragraph, level: int) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    outline = ppr.find(qn("w:outlineLvl"))
    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        ppr.append(outline)
    outline.set(qn("w:val"), str(level))


def _short_title(title: str) -> str:
    title = re.sub(r"[《》]", "", title or "论文题目")
    return title[:24]


def _has_any(text: str, keywords: Iterable[str]) -> bool:
    return any(keyword in text for keyword in keywords)
