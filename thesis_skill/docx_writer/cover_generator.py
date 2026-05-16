from __future__ import annotations

import re
from datetime import date
from typing import Any, Dict

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from thesis_skill.docx_writer.style_manager import body_style


def add_pdf_reference_cover(document, config: Dict[str, Any]) -> None:
    values = cover_values(config)
    for section in document.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.5)

    for _ in range(3):
        document.add_paragraph("")

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("毕业设计说明书")
    run.bold = True
    run.font.size = Pt(26)

    for _ in range(3):
        document.add_paragraph("")

    label = document.add_paragraph(style=body_style(document))
    label.paragraph_format.left_indent = Cm(3.1)
    label.add_run("题目：").bold = True

    topic = document.add_paragraph()
    topic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    topic_run = topic.add_run(values["thesis_title"])
    topic_run.bold = True
    topic_run.font.size = Pt(20)

    for _ in range(2):
        document.add_paragraph("")

    submit = document.add_paragraph(style=body_style(document))
    submit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    submit.add_run(f"提交时间：{values['submit_date']}")

    for _ in range(4):
        document.add_paragraph("")

    for label_text, key in [
        ("姓    名", "student_name"),
        ("班    级", "class_name"),
        ("系    部", "department"),
        ("专    业", "major"),
        ("指导教师", "advisor_name"),
    ]:
        row = document.add_paragraph(style=body_style(document))
        row.paragraph_format.left_indent = Cm(5.2)
        row.paragraph_format.space_after = Pt(10)
        row.add_run(label_text)
        row.add_run("    ")
        row.add_run(values[key])

    document.add_page_break()


def cover_values(config: Dict[str, Any]) -> Dict[str, str]:
    thesis = config.get("thesis", {}) if config else {}
    root = config or {}
    project_name = _first(root.get("project_name"), root.get("project", {}).get("project_name") if isinstance(root.get("project"), dict) else "", thesis.get("project_name"), thesis.get("title"), "食惠零食批发网站")
    values = {
        "student_name": _first(root.get("student_name"), thesis.get("student_name"), thesis.get("author"), "（请填写）"),
        "class_name": _first(root.get("class_name"), thesis.get("class_name"), thesis.get("class"), "（请填写）"),
        "department": _first(root.get("department"), thesis.get("department"), thesis.get("college"), "信息工程系"),
        "major": _first(root.get("major"), thesis.get("major"), "计算机应用技术"),
        "advisor_name": _first(root.get("advisor_name"), thesis.get("advisor_name"), thesis.get("supervisor"), "（请填写）"),
        "project_name": _clean_project_name(project_name),
        "submit_date": _first(root.get("submit_date"), thesis.get("submit_date"), _today_cn()),
    }
    values["thesis_title"] = _first(root.get("thesis_title"), thesis.get("thesis_title"), _format_thesis_title(values["project_name"]))
    values["thesis_title"] = _fix_title_spacing(values["thesis_title"], values["project_name"])
    return values


def _format_thesis_title(project_name: str) -> str:
    clean = _clean_project_name(project_name)
    return f"“{clean}”设计与实现"


def _fix_title_spacing(title: str, project_name: str) -> str:
    clean = re.sub(r"\s+", "", str(title or ""))
    clean = clean.replace("“”", "").replace('" "', "")
    if not clean or clean in {"设计与实现", "网站设计与实现"}:
        return _format_thesis_title(project_name)
    if "“" in clean and "”" in clean:
        return clean
    match = re.match(r"(.+?)(网站)?设计与实现$", clean)
    if match:
        name = match.group(1)
        if match.group(2) and not name.endswith("网站"):
            name += "网站"
        return f"“{_clean_project_name(name)}”设计与实现"
    return _format_thesis_title(project_name)


def _clean_project_name(value: str) -> str:
    text = re.sub(r"[“”\"' ]", "", str(value or "")).strip()
    text = re.sub(r"(的)?设计与实现$", "", text)
    return text or "食惠零食批发网站"


def _first(*values: Any) -> str:
    for value in values:
        if value is None:
            continue
        text = str(value).strip()
        if text:
            return text
    return ""


def _today_cn() -> str:
    today = date.today()
    return f"{today.year}年{today.month}月{today.day}日"
