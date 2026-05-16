from __future__ import annotations

from pathlib import Path

from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
from PIL import Image

from thesis_skill.docx_writer.style_manager import caption_style


def insert_figure_with_caption(
    document: DocxDocument,
    image_path: str | Path,
    caption: str,
    width_cm: float = 13.0,
) -> bool:
    """Insert a centered figure and a caption below it.

    The function keeps the visual style deliberately black-and-white friendly
    and normalizes raster metadata to 300 dpi when Pillow can read the file.
    """

    path = Path(image_path)
    if not path.exists():
        paragraph = document.add_paragraph(f"【图片未找到: {path}】", style=caption_style(document))
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return False

    normalized_path = _ensure_300_dpi(path)
    document.add_picture(str(normalized_path), width=Cm(width_cm))
    image_para = document.paragraphs[-1]
    image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    caption_para = document.add_paragraph(caption, style=caption_style(document))
    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return True


def add_figure(document: DocxDocument, image_path: str | Path, caption: str, width_inches: float = 5.6) -> bool:
    return insert_figure_with_caption(document, image_path, caption, width_cm=width_inches * 2.54)


def _ensure_300_dpi(path: Path) -> Path:
    if path.suffix.lower() not in {".png", ".jpg", ".jpeg"}:
        return path
    try:
        with Image.open(path) as image:
            dpi = image.info.get("dpi", (0, 0))
            if dpi and dpi[0] >= 300 and dpi[1] >= 300:
                return path
            copy_path = path.with_name(f"{path.stem}_300dpi{path.suffix}")
            image.save(copy_path, dpi=(300, 300))
            return copy_path
    except Exception:
        return path
