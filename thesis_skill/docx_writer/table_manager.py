from __future__ import annotations

from typing import Sequence

from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH

from thesis_skill.docx_writer.style_manager import body_style, caption_style


def insert_table_with_caption(document: DocxDocument, caption: str, rows: Sequence[Sequence[str]]) -> None:
    """Insert a caption above a Word table, as required by thesis style."""

    caption_para = document.add_paragraph(caption, style=caption_style(document))
    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if not rows:
        return

    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=column_count)
    table.style = _table_style(document)
    for row_index, row in enumerate(rows):
        for col_index in range(column_count):
            value = row[col_index] if col_index < len(row) else ""
            cell = table.cell(row_index, col_index)
            cell.text = str(value)
            for paragraph in cell.paragraphs:
                paragraph.style = body_style(document)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
    document.add_paragraph("")


def add_data_table(document: DocxDocument, caption: str, rows: Sequence[Sequence[str]]) -> None:
    insert_table_with_caption(document, caption, rows)


def _table_style(document: DocxDocument) -> str:
    names = {style.name for style in document.styles if getattr(style, "type", None)}
    for candidate in ["Table Grid", "网格型", "表格网格"]:
        if candidate in names:
            return candidate
    return "Table Grid"
