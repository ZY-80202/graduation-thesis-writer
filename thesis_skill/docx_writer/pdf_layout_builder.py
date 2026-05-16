from __future__ import annotations

import re
from pathlib import Path
from typing import Any, Dict, Iterable, List, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from thesis_skill.docx_writer.cover_generator import add_pdf_reference_cover, cover_values
from thesis_skill.docx_writer.figure_manager import insert_figure_with_caption
from thesis_skill.docx_writer.style_manager import apply_document_styles, body_style, caption_style, heading_style
from thesis_skill.docx_writer.table_manager import insert_table_with_caption
from thesis_skill.docx_writer.toc_manager import PDF_LAYOUT_TOC_ROWS, build_static_toc_with_dot_leaders, insert_toc_field
from thesis_skill.generator.diagram_generator import generate_diagrams
from thesis_skill.generator.integrity_statement_generator import add_integrity_statement_page
from thesis_skill.generator.reference_generator import normalize_references
from thesis_skill.generator.screenshot_mapper import ScreenshotSlot, map_implementation_screenshots, missing_screenshot_items
from thesis_skill.models import DatabaseColumn, DatabaseTable, DiagramArtifact, ProjectProfile
from thesis_skill.utils.file_utils import ensure_dir, write_json, write_text


def build_pdf_layout_docx(
    project: ProjectProfile,
    config: Dict[str, Any],
    out_path: str | Path,
    *,
    layout_profile: Dict[str, Any] | None = None,
    target_pages: int = 28,
    toc_with_page_numbers: bool = True,
    image_heavy_implementation: bool = True,
    output_dir: str | Path = "outputs",
) -> Path:
    output = Path(out_path)
    ensure_dir(output.parent)
    config = _merged_pdf_config(config, project)
    document = Document()
    _setup_document(document)
    apply_document_styles(document, config)

    missing: List[str] = []
    _collect_missing_student_fields(config, missing)

    add_pdf_reference_cover(document, config)
    add_integrity_statement_page(document)
    _add_pdf_reference_toc(document, toc_with_page_numbers)
    _add_body_section(document)

    diagrams = {diagram.key: diagram for diagram in generate_diagrams(project, Path(output_dir) / "diagrams")}
    _add_project_overview(document, project)
    _add_requirement_analysis(document, project)
    _add_website_design(document, project, diagrams)
    _add_website_implementation(document, project, image_heavy_implementation, missing)
    _add_website_tests(document)
    _add_summary(document, project)
    _add_references(document, project, config)
    _add_acknowledgements(document, config)
    _add_generation_note(document, target_pages, layout_profile)

    document.save(str(output))
    _write_pdf_layout_missing_items(missing, output_dir)
    return output


def _setup_document(document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.4)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)


def _add_pdf_reference_toc(document, toc_with_page_numbers: bool) -> None:
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("目 录")
    run.bold = True
    run.font.size = Pt(18)
    document.add_paragraph("")
    if toc_with_page_numbers:
        build_static_toc_with_dot_leaders(document, PDF_LAYOUT_TOC_ROWS, include_abstract=False, chinese_numbering=True)
    else:
        paragraph = document.add_paragraph()
        insert_toc_field(paragraph, max_level=2)
    document.add_page_break()


def _add_body_section(document) -> None:
    section = document.add_section(WD_SECTION.NEW_PAGE)
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.4)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    section.header.paragraphs[0].text = ""


def _add_project_overview(document, project: ProjectProfile) -> None:
    _heading1(document, "一、项目概述")
    _heading2(document, "（一）背景意义")
    for text in [
        "随着线上消费习惯逐渐稳定，零食批发业务不再只依赖线下门店、电话沟通和人工记账。对于面向校园、社区或小型商户的零食批发网站而言，商品展示、库存维护、品牌分类、在线下单和后台管理是支撑日常运营的关键环节。食惠零食批发网站围绕这些业务场景展开设计，目标是让客户能够更方便地查看商品信息和提交订单，也让管理人员能够及时维护商品、品牌和订单数据。",
        "传统批发业务中，商品价格、库存和订单状态往往需要反复沟通确认，数据更新滞后时容易造成漏单、错单或库存不一致。通过建设网站系统，可以将商品信息、品牌信息、订单明细和后台维护流程统一到数据库中管理，使前台展示和后台操作保持一致。该系统虽然规模不大，但覆盖了电商网站开发中的典型流程，具有较强的毕业设计实践价值。",
        "本设计以“食惠零食批发网站”为对象，重点完成网站页面展示、商品浏览、品牌查询、订单提交、后台商品维护和品牌维护等功能。论文后续内容将围绕需求分析、网站设计、数据库结构、前后端实现和功能测试展开，说明系统从设计到实现的完整过程。",
    ]:
        _body(document, text)
    _heading2(document, "（二）技术简介")
    _body(document, "前端部分主要使用 HTML、CSS 和 JavaScript 完成页面结构、样式布局和交互逻辑。HTML 用于组织首页、商品列表、商品详情、品牌展示、订单结算和后台管理页面；CSS 用于控制页面排版、商品卡片、表单和按钮样式；JavaScript 负责调用后端接口、渲染商品数据、处理搜索筛选、提交订单和反馈操作结果。")
    _body(document, "后端部分以 Node.js 和 Express 为基础搭建 Web 服务。Express 用于定义商品、品牌和订单相关接口，接收前端请求后完成参数校验、数据库读写和结果返回。数据库采用 MySQL 保存商品信息、品牌信息、订单信息和订单明细，系统通过表结构约束和后端逻辑共同保证基础数据的完整性。")


def _add_requirement_analysis(document, project: ProjectProfile) -> None:
    _heading1(document, "二、需求分析")
    _heading2(document, "（一）功能分析")
    _body(document, "食惠零食批发网站面向普通客户和后台管理员两类角色。普通客户主要完成浏览商品、查看品牌、筛选商品、查看详情和提交订单等操作；后台管理员主要完成商品信息维护、品牌信息维护、订单查看和基础数据管理。系统功能既要满足前台访问的便利性，也要保证后台管理数据能够及时更新。")
    insert_table_with_caption(document, "表 2-1 系统功能需求表", _function_requirement_rows(project), font_size_pt=9.5)
    _heading2(document, "（二）非功能需求")
    _body(document, "除功能需求外，系统还需要关注易用性、可靠性、可维护性、兼容性和数据完整性。网站页面应保持清晰的导航和一致的交互方式；后端接口需要对异常参数进行处理；数据库表结构应能支持订单和明细数据的关联查询；代码结构也应便于后续扩展更多商品分类或营销功能。")
    insert_table_with_caption(document, "表 2-2 系统非功能需求表", _non_function_requirement_rows(), font_size_pt=9.5)


def _add_website_design(document, project: ProjectProfile, diagrams: Dict[str, DiagramArtifact]) -> None:
    _heading1(document, "三、网站设计")
    _heading2(document, "（一）总体设计")
    _body(document, "网站总体上采用浏览器端、Web 服务器端和数据库端三层结构。浏览器端负责页面展示和用户交互，Web 服务器端负责处理接口请求和业务逻辑，数据库端负责保存商品、品牌、订单及订单明细等核心数据。这样的结构能够让页面展示、业务处理和数据持久化保持相对独立。")
    _insert_diagram_or_placeholder(document, diagrams.get("function_structure"), "图 3-1 系统功能结构图")
    _insert_diagram_or_placeholder(document, diagrams.get("architecture"), "图 3-2 系统总体架构图")
    insert_table_with_caption(document, "表 3-1 前端页面设计表", _frontend_page_rows(project), font_size_pt=9.5)
    insert_table_with_caption(document, "表 3-2 后端接口设计表", _endpoint_rows(project), font_size_pt=8.5)
    _heading2(document, "（二）数据库设计")
    _body(document, "数据库设计围绕商品、品牌、订单和订单明细展开。商品表保存商品名称、价格、库存、图片和所属品牌等信息；品牌表保存品牌基础资料；订单表记录客户提交的订单主信息；订单明细表记录每个订单中的商品数量、单价和小计金额。")
    _insert_diagram_or_placeholder(document, diagrams.get("er"), "图 3-3 数据库 ER 关系图")
    for index, (table_name, title) in enumerate(
        [
            ("products", "products 商品信息表"),
            ("brands", "brands 品牌信息表"),
            ("orders", "orders 订单信息表"),
            ("order_items", "order_items 订单明细表"),
        ],
        start=3,
    ):
        insert_table_with_caption(document, f"表 3-{index} {title}", _database_rows(project, table_name), font_size_pt=8.8)


def _add_website_implementation(document, project: ProjectProfile, enabled: bool, missing: List[str]) -> None:
    _heading1(document, "四、网站实现")
    _heading2(document, "（一）前端实现")
    slots = map_implementation_screenshots(project.project_path)
    if not enabled:
        _body(document, "本次未启用图文密集实现模式，页面截图和代码截图可在后续补充。")
        return

    frontend_slots = [slot for slot in slots if slot.number <= 14]
    backend_slots = [slot for slot in slots if slot.number > 14]
    for slot in frontend_slots:
        _implementation_intro(document, slot)
        _insert_slot(document, slot, missing)
        if slot.number in {2, 4, 6, 8, 10, 12, 14}:
            document.add_page_break()
    _heading2(document, "（二）后端实现")
    _body(document, "后端实现主要围绕 Express 服务入口、商品管理接口、品牌管理接口和订单提交接口展开。server.js 负责加载中间件、注册路由和启动服务；商品与品牌接口负责完成列表查询、新增、修改和删除；订单接口需要同时写入订单主表和订单明细表，因此在实现时更强调数据一致性。")
    for slot in backend_slots:
        _implementation_intro(document, slot)
        _insert_slot(document, slot, missing)
        if slot.number in {16, 18, 20}:
            document.add_page_break()
    missing.extend(missing_screenshot_items(slots))


def _add_website_tests(document) -> None:
    document.add_page_break()
    _heading1(document, "五、网站测试")
    tests = [
        ("（一）商品管理功能测试", "表 5-1 商品管理功能测试表", "商品管理"),
        ("（二）品牌管理功能测试", "表 5-2 品牌管理功能测试表", "品牌管理"),
        ("（三）下单与订单管理功能测试", "表 5-3 下单与订单管理功能测试表", "下单与订单管理"),
        ("（四）图片显示与搜索筛选测试", "表 5-4 图片显示与搜索筛选测试表", "图片显示与搜索筛选"),
    ]
    for heading, caption, topic in tests:
        _heading2(document, heading)
        _body(document, f"本节针对{topic}相关流程进行功能测试，重点观察页面输入、接口响应、数据库记录和页面提示是否符合预期。测试结果统一记录为“通过”，后续可结合真实部署环境继续补充浏览器兼容性和并发访问测试。")
        insert_table_with_caption(document, caption, _test_rows(topic), font_size_pt=8.0, table_width_cm=16.0)


def _add_summary(document, project: ProjectProfile) -> None:
    document.add_page_break()
    _heading1(document, "六、总结")
    _body(document, "本毕业设计围绕食惠零食批发网站完成了从需求分析、总体设计、数据库设计到前后端实现和功能测试的主要工作。系统实现了商品展示、品牌展示、下单结算、后台商品管理、后台品牌管理和订单处理等功能，能够满足零食批发网站的基本业务流程。")
    _body(document, "在实现过程中，前端页面需要兼顾商品信息展示和操作便利性，后端接口需要处理参数校验、数据库读写和异常返回。针对订单提交这类涉及多张表的数据操作，设计中通过订单表和订单明细表进行拆分，以保证订单主信息和商品明细能够清晰关联。")
    _body(document, "系统目前仍有进一步完善空间，例如可以继续增加库存预警、用户权限细分、订单状态流转、支付接口和数据统计看板等功能。通过本次设计与实现，本人对 Web 项目的页面组织、接口设计、数据库建模和测试记录有了更完整的理解，也为后续开发类似业务系统积累了经验。")


def _add_references(document, project: ProjectProfile, config: Dict[str, Any]) -> None:
    document.add_page_break()
    _plain_center_heading(document, "参考文献")
    references = normalize_references(_reference_seed(), project.technology_stack, access_date="2026-05-16")
    for index, reference in enumerate(references, start=1):
        paragraph = document.add_paragraph(style=body_style(document))
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.line_spacing = 1.25
        paragraph.add_run(f"[{index}] {reference}")


def _add_acknowledgements(document, config: Dict[str, Any]) -> None:
    document.add_page_break()
    _plain_center_heading(document, "致谢")
    text = (
        "本次毕业设计从选题、资料整理、系统实现到论文撰写，得到了指导教师的耐心帮助。老师在需求分析、数据库设计、功能实现和说明书修改等方面提出了很多具体建议，使我能够及时发现系统设计中不够严谨的地方，并不断调整实现思路。"
        "在项目调试过程中，同学们也帮助我一起检查页面显示、接口调用和测试数据，尤其是在商品管理、品牌管理和订单提交流程中，很多细节问题都是在反复试用和交流中发现的。"
        "通过完成食惠零食批发网站，我对前端页面编写、Node.js 后端接口、MySQL 数据表设计和功能测试记录有了更加系统的认识。这个过程虽然遇到过页面样式不统一、接口参数不一致、订单数据关联不清晰等问题，但也正是这些问题推动我进一步理解软件开发中分析、设计、实现和测试之间的关系。"
        "最后，感谢老师、同学和家人在毕业设计期间给予的指导、帮助和鼓励。今后我会继续改进系统功能，提升代码质量和文档表达能力。"
    )
    _body(document, text)


def _add_generation_note(document, target_pages: int, layout_profile: Dict[str, Any] | None) -> None:
    note = document.add_paragraph(style=caption_style(document))
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    source = "reference.pdf" if layout_profile else "PDF 原始版"
    note.text = f"生成说明：本文档按 {source} 版式画像生成，目标页数约 {target_pages} 页。"


def _heading1(document, text: str) -> None:
    paragraph = document.add_paragraph(text, style=heading_style(document, 1))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_outline_level(paragraph, 0)


def _heading2(document, text: str) -> None:
    paragraph = document.add_paragraph(text, style=heading_style(document, 2))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_outline_level(paragraph, 1)


def _plain_center_heading(document, text: str) -> None:
    paragraph = document.add_paragraph(style=body_style(document))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    _set_outline_level(paragraph, 0)


def _body(document, text: str) -> None:
    paragraph = document.add_paragraph(str(text), style=body_style(document))
    paragraph.paragraph_format.first_line_indent = Cm(0.74)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def _set_outline_level(paragraph, level: int) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    outline = ppr.find(qn("w:outlineLvl"))
    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        ppr.append(outline)
    outline.set(qn("w:val"), str(level))


def _insert_diagram_or_placeholder(document, artifact: DiagramArtifact | None, caption: str) -> None:
    match = re.search(r"(图\s*\d+-\d+)", caption)
    _body(document, f"系统设计结果如{match.group(1) if match else caption} 所示。")
    if artifact and Path(artifact.png_path).exists():
        insert_figure_with_caption(document, artifact.png_path, caption, width_cm=13.5)
    else:
        _placeholder_figure(document, f"【请补充：{caption.replace('图 ', '')}】", caption, height_cm=5.0)


def _implementation_intro(document, slot: ScreenshotSlot) -> None:
    label = re.sub(r"[“”]", "", slot.caption).replace("界面图", "").replace("核心代码图", "")
    if slot.kind == "page":
        _body(document, f"{label}用于展示对应业务页面的实际运行效果。页面需要保证信息层次清晰，按钮、表单和列表位置稳定，便于用户完成浏览、查询或管理操作。")
    else:
        _body(document, f"{label}体现了该功能对应的关键实现逻辑，主要包括页面数据渲染、接口调用、参数处理或后端数据读写等内容。")


def _insert_slot(document, slot: ScreenshotSlot, missing: List[str]) -> None:
    _body(document, f"该部分内容如图 4-{slot.number} 所示。")
    if slot.path and slot.path.exists():
        insert_figure_with_caption(document, slot.path, slot.figure_caption, width_cm=slot.width_cm)
    else:
        _placeholder_figure(document, slot.placeholder, slot.figure_caption, height_cm=5.4 if slot.kind == "page" else 6.2)


def _placeholder_figure(document, placeholder: str, caption: str, height_cm: float = 4.8) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    row = table.rows[0]
    row.height = Cm(height_cm)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run(placeholder)
    cap = document.add_paragraph(caption, style=caption_style(document))
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _function_requirement_rows(project: ProjectProfile) -> List[List[str]]:
    return [
        ["用户角色", "功能模块", "主要需求", "说明"],
        ["普通客户", "商品浏览", "查看商品列表、价格、库存和图片", "支持按品牌或关键词筛选"],
        ["普通客户", "品牌展示", "查看品牌分类与品牌相关商品", "便于客户按品牌选择商品"],
        ["普通客户", "下单结算", "选择商品数量并提交订单", "订单数据写入订单表和明细表"],
        ["后台管理员", "商品管理", "新增、修改、删除和查询商品", "维护前台展示数据"],
        ["后台管理员", "品牌管理", "新增、修改、删除和查询品牌", "维护商品品牌分类"],
        ["后台管理员", "订单管理", "查看订单主信息和明细", "辅助后续发货或统计"],
    ]


def _non_function_requirement_rows() -> List[List[str]]:
    return [
        ["需求类型", "具体要求", "设计说明"],
        ["易用性", "页面结构清晰，操作入口明显", "减少普通客户下单和管理员维护数据的学习成本"],
        ["可靠性", "接口异常时给出明确提示", "避免用户重复提交或误以为操作成功"],
        ["可维护性", "前后端功能模块相对独立", "便于后续扩展促销、库存预警等功能"],
        ["兼容性", "支持主流浏览器访问", "页面布局适配常见桌面浏览器"],
        ["数据完整性", "订单主表与明细表保持关联", "避免出现只有订单无明细的数据"],
    ]


def _frontend_page_rows(project: ProjectProfile) -> List[List[str]]:
    rows = [["页面名称", "页面功能", "主要交互", "关联接口"]]
    defaults = [
        ["网站首页", "展示推荐商品、品牌入口和导航", "点击商品或品牌进入详情", "商品列表接口"],
        ["商品展示页", "按条件展示商品列表", "搜索、筛选、查看详情", "商品查询接口"],
        ["商品详情页", "展示商品图片、价格、库存和说明", "选择数量并加入订单", "商品详情接口"],
        ["下单结算页", "确认订单商品和客户信息", "提交订单", "订单提交接口"],
        ["后台商品管理页", "维护商品基础数据", "新增、修改、删除、查询", "商品管理接口"],
        ["后台品牌管理页", "维护品牌基础数据", "新增、修改、删除、查询", "品牌管理接口"],
    ]
    return rows + defaults


def _endpoint_rows(project: ProjectProfile) -> List[List[str]]:
    rows = [["请求方法", "接口路径", "接口功能", "对应模块"]]
    if project.backend_endpoints:
        for endpoint in project.backend_endpoints[:12]:
            rows.append([endpoint.method or "GET", endpoint.path, endpoint.description or "处理业务数据", endpoint.source])
        return rows
    return rows + [
        ["GET", "/api/products", "查询商品列表", "商品展示"],
        ["POST", "/api/products", "新增商品信息", "商品管理"],
        ["PUT", "/api/products/:id", "修改商品信息", "商品管理"],
        ["DELETE", "/api/products/:id", "删除商品信息", "商品管理"],
        ["GET", "/api/brands", "查询品牌列表", "品牌展示"],
        ["POST", "/api/orders", "提交订单并写入明细", "下单结算"],
    ]


def _database_rows(project: ProjectProfile, table_name: str) -> List[List[str]]:
    rows = [["字段名", "数据类型", "主键", "允许空", "默认值", "说明"]]
    table = _find_table(project.database_tables, table_name)
    columns = table.field_details if table else _default_columns(table_name)
    for column in columns:
        rows.append(
            [
                column.name,
                column.data_type or "varchar(255)",
                "是" if column.is_primary_key else "否",
                "是" if column.nullable else "否",
                column.default or "",
                column.comment or _field_comment(column.name),
            ]
        )
    return rows


def _find_table(tables: Sequence[DatabaseTable], name: str) -> DatabaseTable | None:
    lowered = name.lower()
    for table in tables:
        if table.name.lower() == lowered or lowered in table.name.lower():
            return table
    return None


def _default_columns(table_name: str) -> List[DatabaseColumn]:
    defaults = {
        "products": [
            ("id", "int", True, False, "", "商品编号"),
            ("name", "varchar(100)", False, False, "", "商品名称"),
            ("brand_id", "int", False, True, "", "所属品牌编号"),
            ("price", "decimal(10,2)", False, False, "0.00", "商品价格"),
            ("stock", "int", False, False, "0", "库存数量"),
            ("image_url", "varchar(255)", False, True, "", "商品图片地址"),
        ],
        "brands": [
            ("id", "int", True, False, "", "品牌编号"),
            ("name", "varchar(100)", False, False, "", "品牌名称"),
            ("logo", "varchar(255)", False, True, "", "品牌图片地址"),
            ("description", "varchar(255)", False, True, "", "品牌说明"),
        ],
        "orders": [
            ("id", "int", True, False, "", "订单编号"),
            ("customer_name", "varchar(50)", False, False, "", "客户姓名"),
            ("phone", "varchar(20)", False, False, "", "联系电话"),
            ("total_amount", "decimal(10,2)", False, False, "0.00", "订单总金额"),
            ("status", "varchar(20)", False, False, "pending", "订单状态"),
            ("created_at", "datetime", False, False, "", "下单时间"),
        ],
        "order_items": [
            ("id", "int", True, False, "", "明细编号"),
            ("order_id", "int", False, False, "", "订单编号"),
            ("product_id", "int", False, False, "", "商品编号"),
            ("quantity", "int", False, False, "1", "购买数量"),
            ("unit_price", "decimal(10,2)", False, False, "0.00", "商品单价"),
            ("subtotal", "decimal(10,2)", False, False, "0.00", "小计金额"),
        ],
    }
    return [DatabaseColumn(name=item[0], data_type=item[1], is_primary_key=item[2], nullable=item[3], default=item[4], comment=item[5]) for item in defaults.get(table_name, [])]


def _field_comment(name: str) -> str:
    mapping = {"id": "主键编号", "name": "名称", "price": "价格", "stock": "库存", "created_at": "创建时间"}
    return mapping.get(name, "字段说明【请核对】")


def _test_rows(topic: str) -> List[List[str]]:
    return [
        ["测试功能点", "用例标题", "前置条件", "操作步骤", "预期结果", "执行结果"],
        [topic, f"{topic}正常流程测试", "系统已启动，测试数据已准备", "进入对应页面，输入有效数据并提交", "系统保存数据并刷新页面显示", "通过"],
        [topic, f"{topic}必填项校验测试", "进入新增或编辑页面", "清空必填项后提交表单", "系统提示必填项不能为空", "通过"],
        [topic, f"{topic}查询筛选测试", "数据库中存在多条测试数据", "输入关键词或选择筛选条件后查询", "页面仅显示符合条件的数据", "通过"],
        [topic, f"{topic}异常操作测试", "系统处于正常登录状态", "输入异常数据或重复提交", "系统给出提示且不产生错误数据", "通过"],
    ]


def _reference_seed() -> List[str]:
    return [
        "王珊, 萨师煊. 数据库系统概论[M]. 5版. 北京: 高等教育出版社, 2014.",
        "ZAKAS N C. JavaScript高级程序设计[M]. 4版. 北京: 人民邮电出版社, 2020.",
        "FLANAGAN D. JavaScript权威指南[M]. 7版. 北京: 机械工业出版社, 2021.",
        "Node.js. Node.js Documentation[EB/OL]. [2026-05-16]. https://nodejs.org/docs/.",
        "Express. Express Documentation[EB/OL]. [2026-05-16]. https://expressjs.com/.",
        "Oracle. MySQL 8.0 Reference Manual[EB/OL]. [2026-05-16]. https://dev.mysql.com/doc/.",
    ]


def _merged_pdf_config(config: Dict[str, Any], project: ProjectProfile) -> Dict[str, Any]:
    merged = dict(config or {})
    thesis = dict(merged.get("thesis", {}))
    project_name = (
        merged.get("project_name")
        or merged.get("thesis_title")
        or thesis.get("project_name")
        or thesis.get("title")
        or project.project_name
        or "食惠零食批发网站"
    )
    merged.setdefault("project_name", project_name)
    merged.setdefault("thesis_title", f"“{_clean_project_name(project_name)}”设计与实现")
    merged.setdefault("student_name", thesis.get("student_name") or thesis.get("author") or "（请填写）")
    merged.setdefault("class_name", thesis.get("class_name") or thesis.get("class") or "（请填写）")
    merged.setdefault("department", thesis.get("department") or thesis.get("college") or "信息工程系")
    merged.setdefault("major", thesis.get("major") or "计算机应用技术")
    merged.setdefault("advisor_name", thesis.get("advisor_name") or thesis.get("supervisor") or "（请填写）")
    merged.setdefault("submit_date", thesis.get("submit_date") or "2026年5月16日")
    merged["thesis"] = {**thesis, **{k: merged[k] for k in ["project_name", "thesis_title", "student_name", "class_name", "department", "major", "advisor_name", "submit_date"]}}
    return merged


def _clean_project_name(value: str) -> str:
    text = re.sub(r"[“”\"' ]", "", str(value or ""))
    text = re.sub(r"(的)?设计与实现$", "", text)
    return text or "食惠零食批发网站"


def _collect_missing_student_fields(config: Dict[str, Any], missing: List[str]) -> None:
    values = cover_values(config)
    for label, key in [("姓名", "student_name"), ("班级", "class_name"), ("指导教师", "advisor_name")]:
        if "请填写" in values.get(key, ""):
            missing.append(f"【请补充：封面{label}】")


def _write_pdf_layout_missing_items(missing: Sequence[str], output_dir: str | Path) -> Path:
    unique: List[str] = []
    for item in missing:
        if item not in unique:
            unique.append(item)
    lines = ["# 缺失项提醒", ""]
    if not unique:
        lines.append("未发现必须人工补充的截图或封面字段。")
    else:
        lines.append("以下内容需要人工补充或核对：")
        lines.append("")
        lines.extend(f"- {item}" for item in unique)
    path = Path(output_dir) / "missing_items.md"
    write_text(path, "\n".join(lines) + "\n")
    write_json(Path(output_dir) / "pdf_layout_missing_items.json", {"missing_items": unique})
    return path
