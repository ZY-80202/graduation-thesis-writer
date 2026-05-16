from __future__ import annotations

import re
from typing import Iterable, List

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from thesis_skill.docx_writer.numbering_manager import heading_text
from thesis_skill.models import SectionDraft


def add_toc(paragraph: Paragraph) -> None:
    insert_toc_field(paragraph)


def insert_toc_field(paragraph: Paragraph) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = r'TOC \o "1-3" \h \z \u'
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "请在 Word 中更新目录域"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_begin, instr_text, fld_char_separate, placeholder, fld_char_end])


def build_static_toc_with_dot_leaders(document, sections: Iterable[SectionDraft], include_abstract: bool = True) -> None:
    if include_abstract:
        _toc_line(document, "摘    要", 0, "I")
    for section in _flatten_sections(sections):
        if not section.number:
            continue
        title = heading_text(section.number, section.title)
        _toc_line(document, title, max(0, section.level - 1), "待更新")
    _toc_line(document, "总结", 0, "待更新")
    _toc_line(document, "致谢", 0, "待更新")
    _toc_line(document, "参考文献", 0, "待更新")


def validate_toc_has_page_numbers(document) -> List[str]:
    toc_text = _toc_text(document)
    if not toc_text:
        return ["未发现目录内容。"]
    issues: List[str] = []
    if "…" not in toc_text and "." * 4 not in toc_text:
        issues.append("目录缺少点引导符。")
    if not re.search(r"(待更新|[IVX]+|\b\d+\b)", toc_text):
        issues.append("目录缺少页码或页码占位。")
    return issues


def validate_toc_has_numbered_headings(document) -> List[str]:
    toc_text = _toc_text(document)
    if "1 " not in toc_text and "1\t" not in toc_text:
        return ["目录缺少正文一级章节编号。"]
    if not re.search(r"1\.1\s+", toc_text):
        return ["目录缺少二级章节编号。"]
    return []


def _toc_line(document, text: str, level: int, page: str) -> None:
    try:
        from thesis_skill.docx_writer.style_manager import body_style

        paragraph = document.add_paragraph(style=body_style(document))
    except Exception:
        paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    leader_count = max(6, 42 - len(text) * 2 - level * 4)
    paragraph.add_run("    " * level + text)
    paragraph.add_run(" " + "…" * leader_count + " ")
    paragraph.add_run(page)


def _flatten_sections(sections: Iterable[SectionDraft]):
    for section in sections:
        yield section
        yield from _flatten_sections(section.children)


def _toc_text(document) -> str:
    texts = []
    in_toc = False
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text in {"目    录", "目录"}:
            in_toc = True
            continue
        if in_toc and text:
            texts.append(text)
        if in_toc and re.match(r"^1\s+\S+", text):
            continue
    return "\n".join(texts)
