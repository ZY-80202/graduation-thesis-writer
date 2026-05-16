from __future__ import annotations

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def suppress_header_on_cover(section) -> None:
    section.different_first_page_header_footer = True
    for header in [section.header, section.first_page_header]:
        header.is_linked_to_previous = False
        _clear_part(header)


def apply_template_header_footer(section, short_title: str = "") -> None:
    section.header.is_linked_to_previous = False
    header = section.header
    _clear_part(header)
    paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.text = ""
    paragraph.add_run("江苏工程职业技术学院毕业设计（论文）")
    paragraph.add_run("\t")
    paragraph.add_run(short_title or "论文题目")
    _add_bottom_border(paragraph)


def clear_header_footer_links(section) -> None:
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False


def _clear_part(part) -> None:
    for paragraph in part.paragraphs:
        paragraph.text = ""
    for table in part.tables:
        table._element.getparent().remove(table._element)


def _add_bottom_border(paragraph) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    bottom = pbdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        pbdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
