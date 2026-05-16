from __future__ import annotations

from pathlib import Path
from typing import Any, Dict, List

from docx import Document

from thesis_skill.utils.file_utils import require_file
from thesis_skill.utils.text_utils import normalize_whitespace


def read_docx_paragraphs(path: str | Path, include_empty: bool = False) -> List[Dict[str, Any]]:
    file_path = require_file(path, "Word 文档")
    document = Document(str(file_path))
    paragraphs: List[Dict[str, Any]] = []
    for index, paragraph in enumerate(document.paragraphs):
        text = normalize_whitespace(paragraph.text)
        if not include_empty and not text:
            continue
        paragraphs.append(
            {
                "index": index,
                "text": text,
                "style": paragraph.style.name if paragraph.style else "",
                "alignment": str(paragraph.alignment) if paragraph.alignment is not None else "",
            }
        )
    return paragraphs


def read_docx_text(path: str | Path) -> str:
    return "\n".join(item["text"] for item in read_docx_paragraphs(path) if item["text"])


def read_docx_tables(path: str | Path) -> List[List[List[str]]]:
    file_path = require_file(path, "Word 文档")
    document = Document(str(file_path))
    tables: List[List[List[str]]] = []
    for table in document.tables:
        rows: List[List[str]] = []
        for row in table.rows:
            rows.append([normalize_whitespace(cell.text) for cell in row.cells])
        if rows:
            tables.append(rows)
    return tables


def get_heading_paragraphs(path: str | Path) -> List[Dict[str, Any]]:
    headings: List[Dict[str, Any]] = []
    for item in read_docx_paragraphs(path):
        style = item.get("style", "")
        if "Heading" in style or "标题" in style:
            headings.append(item)
    return headings
