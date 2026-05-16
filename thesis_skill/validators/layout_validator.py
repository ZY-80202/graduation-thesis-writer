from __future__ import annotations

import re
from pathlib import Path
from typing import List

from docx import Document

from thesis_skill.models import ValidationIssue


def validate_figure_references(docx_path: str | Path) -> List[ValidationIssue]:
    document = Document(str(docx_path))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    captions = set(re.findall(r"(图\s*\d+\.\d+)", text))
    references = set(re.findall(r"如\s*(图\s*\d+\.\d+)\s*所示", text))
    issues: List[ValidationIssue] = []
    for caption in sorted(captions - references):
        issues.append(ValidationIssue(severity="warning", message=f"{caption} 已插入但正文未明确引用。", location="图题"))
    for reference in sorted(references - captions):
        issues.append(ValidationIssue(severity="warning", message=f"正文引用了 {reference}，但未找到对应图题。", location="正文"))
    return issues


def validate_table_references(docx_path: str | Path) -> List[ValidationIssue]:
    document = Document(str(docx_path))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    captions = set(re.findall(r"(表\s*\d+\.\d+)", text))
    references = set(re.findall(r"如\s*(表\s*\d+\.\d+)\s*所示", text))
    issues: List[ValidationIssue] = []
    for caption in sorted(captions - references):
        issues.append(ValidationIssue(severity="info", message=f"{caption} 已插入，正文可补充引用语句。", location="表题"))
    for reference in sorted(references - captions):
        issues.append(ValidationIssue(severity="warning", message=f"正文引用了 {reference}，但未找到对应表题。", location="正文"))
    return issues
