from __future__ import annotations

import re
import subprocess
from pathlib import Path
from typing import Iterable, List

from docx import Document

from thesis_skill.docx_writer.template_clone import _find_libreoffice
from thesis_skill.models import ValidationIssue
from thesis_skill.utils.file_utils import ensure_dir, write_text


def render_docx_to_pdf(docx_path: str | Path, output_dir: str | Path = "outputs/rendered") -> Path:
    source = Path(docx_path)
    if not source.exists():
        raise FileNotFoundError(f"待渲染 DOCX 不存在: {source}")
    out_dir = ensure_dir(output_dir)
    soffice = _find_libreoffice()
    if not soffice:
        raise RuntimeError("未找到 LibreOffice/soffice，无法执行渲染校验。")
    subprocess.run(
        [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(out_dir), str(source)],
        check=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
    )
    pdf = out_dir / f"{source.stem}.pdf"
    if not pdf.exists():
        raise RuntimeError("LibreOffice 未生成 PDF。")
    return pdf


def render_pdf_to_png(pdf_path: str | Path, output_dir: str | Path = "outputs/rendered/pages", pages: Iterable[int] | None = None) -> List[Path]:
    try:
        import fitz  # type: ignore
    except Exception as exc:  # pragma: no cover
        raise RuntimeError("未安装 PyMuPDF，无法将 PDF 渲染为 PNG。") from exc

    pdf = Path(pdf_path)
    out_dir = ensure_dir(output_dir)
    doc = fitz.open(str(pdf))
    page_indexes = list(pages) if pages is not None else list(range(min(6, doc.page_count)))
    if doc.page_count > 3:
        page_indexes.extend(range(max(0, doc.page_count - 3), doc.page_count))
    page_indexes = sorted(set(index for index in page_indexes if 0 <= index < doc.page_count))
    outputs: List[Path] = []
    for index in page_indexes:
        page = doc[index]
        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
        target = out_dir / f"page_{index + 1:03d}.png"
        pix.save(str(target))
        outputs.append(target)
    doc.close()
    return outputs


def validate_rendered_docx(
    docx_path: str | Path,
    output_dir: str | Path = "outputs",
    max_pages: int = 35,
) -> List[ValidationIssue]:
    issues: List[ValidationIssue] = []
    try:
        pdf = render_docx_to_pdf(docx_path, Path(output_dir) / "rendered")
        render_pdf_to_png(pdf, Path(output_dir) / "rendered" / "pages")
        issues.extend(validate_page_count(pdf, max_pages=max_pages))
    except Exception as exc:
        issues.append(ValidationIssue(severity="warning", message=str(exc), location="渲染"))

    document = Document(str(docx_path))
    issues.extend(validate_cover_visual(document))
    issues.extend(validate_abstract_not_numbered(document))
    issues.extend(validate_toc_visual(document))
    issues.extend(validate_header_footer_by_page(document))
    issues.extend(validate_image_overflow(document))
    write_render_report(issues, output_dir)
    return issues


def validate_cover_visual(document) -> List[ValidationIssue]:
    issues: List[ValidationIssue] = []
    image_count = sum(1 for rel in document.part.rels.values() if "image" in rel.reltype)
    first_text = "\n".join(p.text for p in document.paragraphs[:40])
    if image_count == 0:
        issues.append(ValidationIssue(severity="warning", message="封面或文档未检测到图片，校徽可能没有保留。", location="封面"))
    if "江苏工程职业技术学院毕业设计（论文）" in first_text:
        issues.append(ValidationIssue(severity="error", message="封面前置页疑似出现正文页眉。", location="封面"))
    if "____" not in first_text and "＿" not in first_text and "下划线" not in first_text:
        issues.append(ValidationIssue(severity="warning", message="封面未检测到明显下划线填写区，请人工核对模板克隆结果。", location="封面"))
    return issues


def validate_abstract_not_numbered(document) -> List[ValidationIssue]:
    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        if re.match(r"^\d+\s*摘\s*要$", text):
            return [ValidationIssue(severity="error", message="摘要被错误编号。", location=f"段落 {index}")]
    return []


def validate_toc_visual(document) -> List[ValidationIssue]:
    text = "\n".join(p.text for p in document.paragraphs)
    issues: List[ValidationIssue] = []
    if "目    录" not in text and "目录" not in text:
        issues.append(ValidationIssue(severity="error", message="未发现目录页。", location="目录"))
    if "TOC" not in document._element.xml and "……" not in text and "待更新" not in text:
        issues.append(ValidationIssue(severity="error", message="目录缺少 Word 域或点引导符。", location="目录"))
    return issues


def validate_page_count(pdf_path: str | Path, max_pages: int = 35) -> List[ValidationIssue]:
    try:
        import fitz  # type: ignore

        doc = fitz.open(str(pdf_path))
        count = doc.page_count
        doc.close()
    except Exception:
        return []
    if count > max_pages:
        return [ValidationIssue(severity="error", message=f"渲染页数为 {count} 页，超过最大限制 {max_pages} 页。", location="全文")]
    return []


def validate_header_footer_by_page(document) -> List[ValidationIssue]:
    issues: List[ValidationIssue] = []
    if document.sections and document.sections[0].header.paragraphs:
        header_text = "\n".join(p.text for p in document.sections[0].header.paragraphs)
        if "江苏工程职业技术学院毕业设计（论文）" in header_text:
            issues.append(ValidationIssue(severity="error", message="封面节仍含正文页眉。", location="封面节"))
    if len(document.sections) < 2:
        issues.append(ValidationIssue(severity="warning", message="文档分节数量偏少，可能未正确区分封面、摘要目录和正文。", location="分节"))
    return issues


def validate_image_overflow(document) -> List[ValidationIssue]:
    # Precise visual overflow requires page coordinate extraction from PDF. This
    # lightweight check catches the common DOCX-side problem: inline shapes wider
    # than an A4 text area.
    issues: List[ValidationIssue] = []
    max_width_emu = 14.5 * 360000
    for index, shape in enumerate(document.inline_shapes, start=1):
        if shape.width and shape.width > max_width_emu:
            issues.append(ValidationIssue(severity="warning", message="图片宽度可能超出论文版心。", location=f"图片 {index}"))
    return issues


def write_render_report(issues: List[ValidationIssue], output_dir: str | Path) -> Path:
    ensure_dir(output_dir)
    lines = ["# 渲染视觉检查报告", ""]
    if not issues:
        lines.append("渲染检查未发现阻断问题。")
    else:
        lines.append(f"发现 {len(issues)} 项问题。")
        lines.append("")
        for index, issue in enumerate(issues, start=1):
            loc = f"（{issue.location}）" if issue.location else ""
            lines.append(f"{index}. [{issue.severity}] {issue.message}{loc}")
    path = Path(output_dir) / "render_check_report.md"
    write_text(path, "\n".join(lines) + "\n")
    return path
