from __future__ import annotations

import re
from pathlib import Path
from typing import List

from docx import Document

from thesis_skill.docx_writer.style_manager import validate_style_usage
from thesis_skill.docx_writer.toc_manager import validate_toc_has_numbered_headings, validate_toc_has_page_numbers
from thesis_skill.models import ValidationIssue, ValidationReport
from thesis_skill.utils.file_utils import ensure_dir, write_json, write_text
from thesis_skill.utils.text_utils import has_placeholder, normalize_whitespace
from thesis_skill.validators.layout_validator import validate_figure_references, validate_table_references
from thesis_skill.validators.reference_validator import extract_references_from_docx, validate_references


def validate_docx(
    docx_path: str | Path,
    template_path: str | Path | None = None,
    output_dir: str | Path = "outputs",
) -> ValidationReport:
    path = Path(docx_path)
    if not path.exists():
        raise FileNotFoundError(f"待检查 Word 不存在: {path}")
    document = Document(str(path))
    issues: List[ValidationIssue] = []
    paragraphs = [(index, normalize_whitespace(p.text), p.style.name if p.style else "") for index, p in enumerate(document.paragraphs)]

    issues.extend(_check_front_matter(paragraphs))
    issues.extend(_check_headings(paragraphs))
    issues.extend(_check_placeholders(paragraphs))
    issues.extend(_check_captions(paragraphs))
    issues.extend(_check_tables(document, paragraphs))
    issues.extend(_check_images(document, paragraphs))
    issues.extend(validate_style_usage(str(path)))
    issues.extend(_toc_issues(document))
    issues.extend(validate_figure_references(path))
    issues.extend(validate_table_references(path))

    ref_issues = validate_references(extract_references_from_docx(path), output_dir)
    issues.extend(ValidationIssue(severity="warning", message=item, location="参考文献") for item in ref_issues)

    summary = "通过基础格式检查。" if not _blocking(issues) else f"发现 {len(issues)} 个需要处理的问题。"
    report = ValidationReport(docx_path=str(path), issues=issues, summary=summary)
    ensure_dir(output_dir)
    write_json(Path(output_dir) / "format_check_report.json", report)
    write_text(Path(output_dir) / "format_check_report.md", _format_markdown(report))
    return report


def _check_front_matter(paragraphs) -> List[ValidationIssue]:
    issues: List[ValidationIssue] = []
    for index, text, _style in paragraphs[:120]:
        if re.match(r"^1\s*摘\s*要$", text):
            issues.append(ValidationIssue(severity="error", message="摘要被错误编号为正文一级标题。", location=f"段落 {index}"))
        if re.match(r"^1\s*目\s*录$", text):
            issues.append(ValidationIssue(severity="error", message="目录被错误编号为正文一级标题。", location=f"段落 {index}"))
    return issues


def _check_headings(paragraphs) -> List[ValidationIssue]:
    issues: List[ValidationIssue] = []
    last_numbers: dict[int, int] = {}
    saw_first_chapter = False
    for index, text, style in paragraphs:
        if not _is_heading_style(style) and not re.match(r"^\d+(\.\d+)*\s+\S+", text):
            continue
        if not text:
            issues.append(ValidationIssue(severity="error", message="存在空标题。", location=f"段落 {index}"))
            continue
        if re.match(r"^\d+\s+(致谢|参考文献|摘\s*要|目\s*录)", text):
            issues.append(ValidationIssue(severity="error", message=f"前置或后置标题不应参与正文编号: {text}", location=f"段落 {index}"))
        level1 = re.match(r"^(\d+)\s+(.+)", text)
        if level1 and "." not in level1.group(1):
            number = int(level1.group(1))
            title = level1.group(2)
            if not saw_first_chapter:
                saw_first_chapter = True
                if number != 1 or "概述" not in title:
                    issues.append(ValidationIssue(severity="error", message="正文第一章应为“1 概述”。", location=f"段落 {index}"))
        match = re.match(r"^(\d+)\.(\d+)", text)
        if match:
            chapter = int(match.group(1))
            number = int(match.group(2))
            last = last_numbers.get(chapter, 0)
            if number != last + 1 and number != 1:
                issues.append(ValidationIssue(severity="warning", message=f"章节编号可能跳跃: {text}", location=f"段落 {index}"))
            last_numbers[chapter] = number
    return issues


def _check_placeholders(paragraphs) -> List[ValidationIssue]:
    issues: List[ValidationIssue] = []
    for index, text, _ in paragraphs:
        if has_placeholder(text):
            issues.append(ValidationIssue(severity="warning", message=f"存在未替换占位符: {text[:60]}", location=f"段落 {index}"))
    return issues


def _check_captions(paragraphs) -> List[ValidationIssue]:
    issues: List[ValidationIssue] = []
    figure_seen: dict[int, int] = {}
    table_seen: dict[int, int] = {}
    for index, text, _ in paragraphs:
        fig = re.match(r"图\s*(\d+)\.(\d+)", text)
        tab = re.match(r"表\s*(\d+)\.(\d+)", text)
        if fig:
            _check_sequence("图题", fig, figure_seen, text, index, issues)
        if tab:
            _check_sequence("表题", tab, table_seen, text, index, issues)
    return issues


def _check_sequence(kind: str, match, seen: dict[int, int], text: str, index: int, issues: List[ValidationIssue]) -> None:
    chapter, number = int(match.group(1)), int(match.group(2))
    expected = seen.get(chapter, 0) + 1
    if number != expected:
        issues.append(ValidationIssue(severity="warning", message=f"{kind}编号不连续: {text}", location=f"段落 {index}"))
    seen[chapter] = max(seen.get(chapter, 0), number)


def _check_tables(document, paragraphs) -> List[ValidationIssue]:
    captions = [text for _, text, _ in paragraphs if re.match(r"表\s*\d+\.\d+", text)]
    if len(captions) < len(document.tables):
        return [ValidationIssue(severity="warning", message="存在表格标题缺失或表题未按“表 x.y”编号。", location="表格")]
    return []


def _check_images(document, paragraphs) -> List[ValidationIssue]:
    image_count = sum(1 for rel in document.part.rels.values() if "image" in rel.reltype)
    figure_captions = sum(1 for _, text, _ in paragraphs if re.match(r"图\s*\d+\.\d+", text))
    if figure_captions and image_count < figure_captions:
        return [ValidationIssue(severity="warning", message="图题数量多于图片数量，可能存在图片未插入。", location="图片")]
    return []


def _toc_issues(document) -> List[ValidationIssue]:
    issues = []
    for message in validate_toc_has_page_numbers(document):
        issues.append(ValidationIssue(severity="error", message=message, location="目录"))
    for message in validate_toc_has_numbered_headings(document):
        issues.append(ValidationIssue(severity="error", message=message, location="目录"))
    return issues


def _is_heading_style(style: str) -> bool:
    return "Heading" in style or "标题" in style


def _blocking(issues: List[ValidationIssue]) -> bool:
    return any(issue.severity == "error" for issue in issues)


def _format_markdown(report: ValidationReport) -> str:
    lines = ["# 文档格式检查报告", "", report.summary, ""]
    if not report.issues:
        return "\n".join(lines) + "\n"
    for index, issue in enumerate(report.issues, start=1):
        location = f"（{issue.location}）" if issue.location else ""
        lines.append(f"{index}. [{issue.severity}] {issue.message}{location}")
    return "\n".join(lines) + "\n"
