from __future__ import annotations

from pathlib import Path
from typing import Dict, List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from thesis_skill.models import PageSettings, ParagraphStyleProfile, TemplateProfile, TemplateRegion
from thesis_skill.utils.file_utils import ensure_dir, write_json
from thesis_skill.utils.text_utils import normalize_whitespace

REGION_KEYWORDS = {
    "cover": ["毕业设计", "题目", "学院", "专业", "指导教师"],
    "abstract_cn": ["摘要"],
    "abstract_en": ["Abstract", "ABSTRACT"],
    "toc": ["目录"],
    "body": ["绪论", "第1章", "第一章"],
    "references": ["参考文献"],
    "acknowledgements": ["致谢"],
    "appendix": ["附录"],
}


def analyze_template(template_path: str | Path, output_dir: str | Path = "outputs/profiles") -> TemplateProfile:
    path = Path(template_path)
    if not path.exists():
        raise FileNotFoundError(f"学校模板不存在: {path}")
    document = Document(str(path))
    page_settings = [_extract_page_settings(section) for section in document.sections]
    styles = _extract_styles(document)
    regions, detected = _detect_regions(document)
    notes: List[str] = []
    if not styles:
        notes.append("未提取到可用样式，生成文档会使用内置默认样式。")
    if not regions:
        notes.append("未识别到封面/摘要/目录等区域关键词，请人工检查模板。")
    profile = TemplateProfile(
        template_path=str(path),
        page_settings=page_settings,
        styles=styles,
        regions=regions,
        detected_keywords=detected,
        notes=notes,
    )
    ensure_dir(output_dir)
    write_json(Path(output_dir) / "template_profile.json", profile)
    return profile


def _extract_page_settings(section) -> PageSettings:
    return PageSettings(
        width_cm=_emu_to_cm(section.page_width),
        height_cm=_emu_to_cm(section.page_height),
        top_margin_cm=_emu_to_cm(section.top_margin),
        bottom_margin_cm=_emu_to_cm(section.bottom_margin),
        left_margin_cm=_emu_to_cm(section.left_margin),
        right_margin_cm=_emu_to_cm(section.right_margin),
        header_distance_cm=_emu_to_cm(section.header_distance),
        footer_distance_cm=_emu_to_cm(section.footer_distance),
        header_text=normalize_whitespace("\n".join(p.text for p in section.header.paragraphs)),
        footer_text=normalize_whitespace("\n".join(p.text for p in section.footer.paragraphs)),
    )


def _extract_styles(document: Document) -> Dict[str, ParagraphStyleProfile]:
    wanted = {
        "Normal",
        "Body Text",
        "Title",
        "Subtitle",
        "Heading 1",
        "Heading 2",
        "Heading 3",
        "正文",
        "标题",
        "标题 1",
        "标题 2",
        "标题 3",
        "图题",
        "表题",
        "参考文献",
    }
    styles: Dict[str, ParagraphStyleProfile] = {}
    for style in document.styles:
        if style.type != 1:
            continue
        if style.name not in wanted and not any(key in style.name for key in ["标题", "Heading", "图", "表", "参考"]):
            continue
        font = style.font
        paragraph_format = style.paragraph_format
        styles[style.name] = ParagraphStyleProfile(
            style_name=style.name,
            font_name=font.name,
            font_size_pt=_emu_to_pt(font.size),
            bold=font.bold,
            italic=font.italic,
            alignment=_alignment_name(paragraph_format.alignment),
            line_spacing=str(paragraph_format.line_spacing) if paragraph_format.line_spacing else None,
            first_line_indent_pt=_emu_to_pt(paragraph_format.first_line_indent),
            space_before_pt=_emu_to_pt(paragraph_format.space_before),
            space_after_pt=_emu_to_pt(paragraph_format.space_after),
        )
    return styles


def _detect_regions(document: Document) -> tuple[List[TemplateRegion], List[str]]:
    regions: List[TemplateRegion] = []
    detected: List[str] = []
    seen_names: set[str] = set()
    for index, paragraph in enumerate(document.paragraphs):
        text = normalize_whitespace(paragraph.text)
        if not text:
            continue
        for name, keywords in REGION_KEYWORDS.items():
            if name in seen_names:
                continue
            for keyword in keywords:
                if keyword in text:
                    regions.append(TemplateRegion(name=name, keyword=keyword, paragraph_index=index))
                    detected.append(keyword)
                    seen_names.add(name)
                    break
    return regions, detected


def _emu_to_cm(value) -> float | None:
    if value is None:
        return None
    return round(float(value.cm), 2)


def _emu_to_pt(value) -> float | None:
    if value is None:
        return None
    return round(float(value.pt), 2)


def _alignment_name(value) -> str | None:
    if value is None:
        return None
    mapping = {
        WD_ALIGN_PARAGRAPH.LEFT: "left",
        WD_ALIGN_PARAGRAPH.CENTER: "center",
        WD_ALIGN_PARAGRAPH.RIGHT: "right",
        WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
    }
    return mapping.get(value, str(value))
