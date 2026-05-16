from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable, List, Sequence

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from thesis_skill.docx_writer.numbering_manager import heading_text
from thesis_skill.models import SectionDraft

PDF_LAYOUT_TOC_ROWS = [
    ("一、项目概述", 1, "1"),
    ("（一）背景意义", 2, "1"),
    ("（二）技术简介", 2, "1"),
    ("二、需求分析", 1, "2"),
    ("（一）功能分析", 2, "2"),
    ("（二）非功能需求", 2, "2"),
    ("三、网站设计", 1, "3"),
    ("（一）总体设计", 2, "3"),
    ("（二）数据库设计", 2, "3"),
    ("四、网站实现", 1, "4"),
    ("（一）前端实现", 2, "4"),
    ("（二）后端实现", 2, "16"),
    ("五、网站测试", 1, "23"),
    ("（一）商品管理功能测试", 2, "23"),
    ("（二）品牌管理功能测试", 2, "24"),
    ("（三）下单与订单管理功能测试", 2, "25"),
    ("（四）图片显示与搜索筛选测试", 2, "25"),
    ("六、总结", 1, "26"),
    ("参考文献", 1, "27"),
    ("致谢", 1, "28"),
]


def add_toc(paragraph: Paragraph) -> None:
    insert_toc_field(paragraph)


def insert_toc_field(paragraph: Paragraph, max_level: int = 2) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = rf'TOC \o "1-{max_level}" \h \z \u'
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "请在 Word 中更新目录域"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_begin, instr_text, fld_char_separate, placeholder, fld_char_end])


def build_static_toc_with_dot_leaders(
    document,
    sections: Iterable[SectionDraft] | Sequence[tuple[str, int, str]] | None = None,
    include_abstract: bool = True,
    *,
    chinese_numbering: bool = False,
    page_numbers: dict[str, str] | None = None,
) -> None:
    if chinese_numbering:
        rows = list(sections or PDF_LAYOUT_TOC_ROWS)  # type: ignore[arg-type]
        for title, level, page in rows:
            _toc_line(document, str(title), int(level), (page_numbers or {}).get(str(title), str(page)), dot=".")
        return

    if include_abstract:
        _toc_line(document, "摘    要", 1, "I", dot="…")
    for section in _flatten_sections(sections or []):  # type: ignore[arg-type]
        if not section.number:
            continue
        title = heading_text(section.number, section.title)
        _toc_line(document, title, max(1, section.level), "待更新", dot="…")
    _toc_line(document, "总结", 1, "待更新", dot="…")
    _toc_line(document, "致谢", 1, "待更新", dot="…")
    _toc_line(document, "参考文献", 1, "待更新", dot="…")


def update_static_toc_page_numbers_from_rendered_pdf(docx_path: str | Path, rendered_pdf: str | Path, toc_rows=None) -> dict[str, str]:
    """Infer heading pages from a rendered PDF.

    Updating an existing DOCX in-place reliably requires a second Word write
    pass. This function returns the inferred mapping so callers can rebuild the
    static TOC with precise page numbers when needed.
    """

    try:
        import fitz  # type: ignore
    except Exception:
        return {}
    rows = toc_rows or PDF_LAYOUT_TOC_ROWS
    targets = [row[0] for row in rows]
    result: dict[str, str] = {}
    doc = fitz.open(str(rendered_pdf))
    for page_index, page in enumerate(doc, start=1):
        text = page.get_text("text")
        for target in targets:
            if target in text and target not in result:
                result[target] = str(page_index)
    doc.close()
    return result


def validate_toc_has_dot_leaders(document) -> List[str]:
    toc_text = _toc_text(document)
    if not toc_text:
        return ["未发现目录内容。"]
    if "." * 6 not in toc_text and "…" not in toc_text:
        return ["目录缺少点引导符。"]
    return []


def validate_toc_has_page_numbers(document) -> List[str]:
    toc_text = _toc_text(document)
    if not toc_text:
        return ["未发现目录内容。"]
    issues: List[str] = []
    issues.extend(validate_toc_has_dot_leaders(document))
    lines = [line.strip() for line in toc_text.splitlines() if line.strip()]
    if lines and not all(re.search(r"(\d+|[IVX]+|待更新)\s*$", line) for line in lines[: min(8, len(lines))]):
        issues.append("目录缺少页码或页码占位。")
    return issues


def validate_toc_has_numbered_headings(document) -> List[str]:
    toc_text = _toc_text(document)
    if "一、项目概述" in toc_text and "（一）背景意义" in toc_text:
        return []
    if "1 " in toc_text and re.search(r"1\.1\s+", toc_text):
        return []
    return ["目录缺少一级/二级章节编号。"]


def _toc_line(document, text: str, level: int, page: str, dot: str = ".") -> None:
    try:
        from thesis_skill.docx_writer.style_manager import body_style

        paragraph = document.add_paragraph(style=body_style(document))
    except Exception:
        paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    indent = "    " if level >= 2 else ""
    leader_count = max(8, 68 - len(text) * 2 - (4 if level >= 2 else 0))
    paragraph.add_run(indent + text)
    paragraph.add_run(dot * leader_count)
    paragraph.add_run(str(page))


def _flatten_sections(sections: Iterable[SectionDraft]):
    for section in sections:
        yield section
        yield from _flatten_sections(section.children)


def _toc_text(document) -> str:
    texts = []
    in_toc = False
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text in {"目    录", "目 录", "目录"}:
            in_toc = True
            continue
        if in_toc and text:
            if re.match(r"^(一、|二、|三、|四、|五、|六、|七、|参考文献|致谢|（)", text) or "." * 4 in text or "…" in text:
                texts.append(text)
                continue
            if re.match(r"^(摘|1\s+|1\.1)", text):
                texts.append(text)
                continue
            if texts:
                break
    return "\n".join(texts)
