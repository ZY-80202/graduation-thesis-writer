from __future__ import annotations

from typing import Sequence

from docx.document import Document as DocxDocument
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from thesis_skill.docx_writer.style_manager import body_style, caption_style


def insert_table_with_caption(
    document: DocxDocument,
    caption: str,
    rows: Sequence[Sequence[str]],
    *,
    font_size_pt: float = 10.5,
    table_width_cm: float = 15.8,
    repeat_header: bool = True,
) -> None:
    """Insert a centered thesis table with caption above it."""

    caption_para = document.add_paragraph(caption, style=caption_style(document))
    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if not rows:
        return

    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=column_count)
    table.style = _table_style(document)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    width = Cm(table_width_cm / max(column_count, 1))

    for row_index, row in enumerate(rows):
        if row_index == 0 and repeat_header:
            _repeat_table_header(table.rows[row_index])
        for col_index in range(column_count):
            value = row[col_index] if col_index < len(row) else ""
            cell = table.cell(row_index, col_index)
            cell.width = width
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = str(value)
            for paragraph in cell.paragraphs:
                paragraph.style = body_style(document)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    run.font.size = Pt(font_size_pt)
                    if row_index == 0:
                        run.bold = True
    document.add_paragraph("")


def add_data_table(document: DocxDocument, caption: str, rows: Sequence[Sequence[str]]) -> None:
    insert_table_with_caption(document, caption, rows)


def _repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def _table_style(document: DocxDocument) -> str:
    names = {style.name for style in document.styles if getattr(style, "type", None)}
    for candidate in ["Table Grid", "网格型", "表格网格"]:
        if candidate in names:
            return candidate
    return "Table Grid"
