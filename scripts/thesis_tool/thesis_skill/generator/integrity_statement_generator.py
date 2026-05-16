from __future__ import annotations

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from thesis_skill.docx_writer.style_manager import body_style


INTEGRITY_TEXT = (
    "本人郑重声明：本毕业设计说明书是在指导教师指导下，结合本人对项目系统的分析、设计、实现和测试过程独立完成的。"
    "文中引用的资料和技术文献均在参考文献中列出，未抄袭、剽窃他人成果。如有不实之处，本人愿意承担相应责任。"
)


def add_integrity_statement_page(document) -> None:
    for _ in range(2):
        document.add_paragraph("")
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("毕业设计诚信声明")
    run.bold = True
    run.font.size = Pt(18)

    document.add_paragraph("")
    paragraph = document.add_paragraph(INTEGRITY_TEXT, style=body_style(document))
    paragraph.paragraph_format.first_line_indent = Cm(0.74)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for _ in range(8):
        document.add_paragraph("")
    sign = document.add_paragraph(style=body_style(document))
    sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sign.add_run("学生签名：________________    日期：________年____月____日")
    document.add_page_break()
